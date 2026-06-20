from __future__ import annotations

import json
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path.cwd()
VERSION = "2026-06-20"
UPDATED_DE = "20. Juni 2026"
OUT_NAME = "VYVA_Admin_Console_User_Manual_DE"
MANUAL_URL = "https://rcadmin.vyva.life/manuals/VYVA_Admin_Console_User_Manual_DE.pdf"

CURRENT_DIR = ROOT / "docs" / "admin-manual" / "current"
VERSION_DIR = ROOT / "docs" / "admin-manual" / "versions" / VERSION
SOURCE_DIR = ROOT / "docs" / "admin-manual" / "source"
PUBLIC_DIR = ROOT / "public" / "manuals"
ARCHIVE_DIR = PUBLIC_DIR / "archive"
SCREENSHOT_DIR = CURRENT_DIR / "screenshots-from-english"
ENGLISH_MANUAL = ROOT / "VYVA_Admin_Console_User_Manual.docx"


SECTIONS = [
    {
        "title": "1. Zweck und Umfang der Konsole",
        "paras": [
            "Die VYVA x Red Cross Konsole ist ein operativer Arbeitsbereich für Betreuung, Nachverfolgung, Risikosignale, Medikationsroutinen, Kampagnen und Versorgungsabdeckung älterer Klientinnen und Klienten.",
            "Dieses Handbuch richtet sich an Administratoren, Koordinatoren, Operatoren und autorisiertes professionelles Personal. Es beschreibt, wie die aktuelle Konsole genutzt wird, welche Daten jede Ansicht liefern soll und welche Aktionen eine externe Integration benötigen.",
            "Die Konsole folgt dem Prinzip der minimal notwendigen Information: Erfassen Sie nur Daten, die für Koordination, Sicherheit und Betreuung wirklich gebraucht werden.",
        ],
        "bullets": [
            "Clients sind die betreuten Personen.",
            "Team members sind interne Benutzerkonten für die Konsole.",
            "Emergency contacts sind Angehörige, Nachbarn oder informelle Betreuungspersonen.",
            "Red Cross staff ist professionelles Personal der Organisation.",
        ],
    },
    {
        "title": "2. Anmeldung",
        "paras": [
            "Die sichtbare Anmeldung erfolgt ausschließlich per E-Mail. Admins geben ihre E-Mail-Adresse ein und erhalten einen sicheren Magic Link.",
            "Google- und Microsoft-OAuth bleiben ausgeblendet, bis diese Anbieter bewusst aktiviert werden. Supabase Auth dient nur als Identitätsschicht für den E-Mail-Link; Replit verwaltet Organisationen, Rollen, Teams und operative Daten.",
            "Nach dem Login ordnet das System die E-Mail einer Rolle und Organisation zu. Unbekannte E-Mails erhalten keinen Konsolenzugriff.",
        ],
        "steps": [
            "Öffnen Sie die veröffentlichte Konsolen-URL.",
            "Geben Sie Ihre Admin-E-Mail-Adresse ein.",
            "Öffnen Sie den neuesten Magic Link aus Ihrem Postfach.",
            "Prüfen Sie nach dem Login die aktive Organisation.",
        ],
    },
    {
        "title": "3. Rollen und Berechtigungen",
        "paras": [
            "Berechtigungen werden aus der authentifizierten E-Mail und den Backend-Profilen abgeleitet. Die Oberfläche soll die echte Rolle anzeigen, zum Beispiel Super admin, Admin, Coordinator oder Operator.",
            "Superadmins werden nicht in der UI erstellt. Sie werden über Backend-Kontrollen wie VYVA_PLATFORM_ADMIN_EMAILS oder Profileinstellungen vergeben.",
        ],
        "table": [
            ["Rolle", "Hauptumfang", "Hinweis"],
            ["Super admin", "Kann Organisationen wechseln und Organisationsadmins erstellen.", "Nicht über die normale UI verwaltet."],
            ["Admin", "Verwaltet Clients, Team, Kampagnen, Kontakte, Zeitpläne und Einstellungen der Organisation.", "Kann Teammitglieder anlegen."],
            ["Coordinator / Operator", "Bearbeitet Queues, Profile und zugewiesene Aufgaben.", "Sensible Aktionen können eingeschränkt sein."],
            ["Primär zugewiesenes Red Cross staff", "Kann bestimmte Care-Plan-Daten bearbeiten, wenn Zustimmung vorliegt.", "Nicht dasselbe wie Emergency contact."],
        ],
    },
    {
        "title": "4. Aktive Organisation",
        "paras": [
            "Die aktive Organisation bestimmt, welche Daten in der Konsole angezeigt werden. Beim Wechsel zwischen Red Cross Zamora und Red Cross Leipzig müssen Clients, Kampagnen, Check-ins, Brain Coach, Medikation, Risiko, Kontakte, Personal, Reports und Karte neu geladen werden.",
            "Red Cross Zamora nutzt standardmäßig Spanisch, Spanien und Europe/Madrid. Red Cross Leipzig nutzt Deutsch, Deutschland und Europe/Berlin.",
            "Die Zuordnung ist strikt: spanische Nummern und Adressen gehören zu Zamora; deutsche Nummern und Adressen gehören zu Leipzig. Es darf kein generischer Fallback Clients in die falsche Organisation verschieben.",
        ],
    },
    {
        "title": "5. Allgemeine Navigation",
        "paras": [
            "Die Seitenleiste gruppiert die Arbeit in Main, Follow-up und Management. Die wichtigsten Bereiche sind Today, Clients, Risk, Sensors, Check-ins, Medication, Brain Coach, Campaigns, Care Providers, Reports, Team access und Settings.",
            "Der Kopfbereich zeigt Konsole, aktive Organisation, Systemstatus und den angemeldeten Benutzer. Bei Superadmins erscheint zusätzlich der Organisationswechsel.",
        ],
        "table": [
            ["Bereich", "Zweck"],
            ["Today", "Tägliche operative Übersicht mit Kennzahlen und Karte."],
            ["Clients", "Client-Queue und Pflegeprofile."],
            ["Risk", "Priorisierte Fälle nach Risiko, Nichtantwort, Medikation oder fehlender Abdeckung."],
            ["Sensors", "Vorbereiteter Bereich für Gerätedaten."],
            ["Check-ins", "Routine-Check-up-Anrufe und Status."],
            ["Brain Coach", "Kognitive Routinen und Aktivitätsberichte."],
            ["Campaigns", "VYVA-Anrufkampagnen."],
        ],
    },
    {
        "title": "6. Today Dashboard",
        "paras": [
            "Today fasst den operativen Zustand der aktiven Organisation zusammen. Kennzahlen müssen eine klare Berechnungslogik und einen eindeutigen Zeitraum haben.",
            "Die Karte steht oben und behält das Leaflet-Verhalten bei: stabile Größe, korrektes Resize, Tiles, Clustering und passende Zentrierung für das Land der aktiven Organisation.",
            "Check-ins werden wöchentlich bewertet: erledigte Check-ins im Verhältnis zu erwarteten Check-ins dieser Woche.",
        ],
        "bullets": [
            "Active clients: aktive Clients in der Organisation.",
            "Urgent: Fälle mit sofortigem Handlungsbedarf.",
            "Review: Fälle für heutige Operator-Prüfung.",
            "Check-ins: wöchentliche Erfüllung.",
            "Medication: offene oder unbestätigte Medikationssignale.",
        ],
    },
    {
        "title": "7. Clients",
        "paras": [
            "Clients ist die Hauptliste betreuter Personen. Sie können nach Name, Telefon, Stadt oder Kontakt suchen und nach Dringlichkeit, Review, No response, Medication, Check-ins oder Unassigned filtern.",
            "Die Aktionen Add client, Import clients und API intake gehören zur Client-Erstellung. Teamkonten werden nicht hier erstellt.",
        ],
        "steps": [
            "Add client für einzelne neue Clients.",
            "Import clients für CSV-Upload.",
            "API intake für Clients, die bereits im Onboarding-Backend existieren.",
            "Eine Zeile öffnen, um Profil und Care coverage zu prüfen.",
        ],
    },
    {
        "title": "8. Clients erstellen oder importieren",
        "paras": [
            "Das Client-Formular ist als Care Profile Intake aufgebaut: Person und Kontakt, medizinisches Mindestprofil, Medikation, Zustimmung, Emergency contact und Follow-up-Routinen.",
            "Telefonnummern müssen im internationalen Format mit + und Ländercode erfasst werden. Das unterstützt Routing, Organisationserkennung und spätere Anrufintegration.",
            "CSV-Importe müssen dieselben Regeln für minimale notwendige Daten beachten. Emergency contacts aus CSV oder Onboarding bleiben persönliche Kontakte und werden nicht zu Red Cross staff.",
        ],
    },
    {
        "title": "9. Client-Profil",
        "paras": [
            "Das Client-Profil ist die zentrale Arbeitsansicht. Es zeigt Name, Stadt, Status, Telefon, Sprache, Adresse, Zustimmung, letzten Kontakt und Care coverage.",
            "Die Karte Medication and check-ins zeigt Medikamente, Zeiten, Frequenz und Routinen. View adherence öffnet den wöchentlichen Medikationskalender des Clients.",
            "Call now, Send WhatsApp und Contact care provider dürfen keine Aktion vortäuschen. Wenn kein Gateway verbunden ist, muss die UI erklären, dass eine Gateway-Verbindung erforderlich ist.",
        ],
    },
    {
        "title": "10. Emergency contacts",
        "paras": [
            "Emergency contacts sind nicht-professionelle Unterstützungspersonen: Familie, Nachbarn, informelle Betreuung oder persönliche Notfallkontakte.",
            "Diese Kontakte werden während Onboarding, eingehenden Gesprächen oder im Admin-Intake erfasst. Die Seite darf kein Red Cross staff enthalten.",
            "Ein Client kann mehrere Emergency contacts haben, und ein Kontakt kann mehreren Clients zugeordnet sein.",
        ],
    },
    {
        "title": "11. Red Cross staff",
        "paras": [
            "Red Cross staff ist professionelles Personal der aktiven Organisation. Es wird getrennt von Emergency contacts verwaltet.",
            "Beim Zuweisen von Personal wird eine professionelle Rolle ausgewählt, nicht eine familiäre Beziehung. Beispiele sind Field coordinator, Primary operator, Medication support, Social worker oder Supervisor.",
        ],
    },
    {
        "title": "12. Check-ins",
        "paras": [
            "Check-ins ist nur für Check-up-Anrufe. Brain Coach Sessions werden auf der separaten Brain Coach Seite geführt.",
            "Jede Zeile zeigt Client, Telefon, Typ, Status, letzten Check-in, Frequenz, bevorzugte Uhrzeit und Aktionen.",
            "Wenn die geplante Uhrzeit vorbei ist und kein erfolgreicher Verlauf existiert, muss die Zeile Missed today anzeigen. Erfolgreiche Anrufe werden als Confirmed oder Completed angezeigt. Eskalationen werden als Escalated markiert.",
        ],
        "table": [
            ["Status", "Bedeutung"],
            ["Active", "Routine ist aktiviert."],
            ["Missed", "Geplanter Anruf ist ohne Bestätigung verstrichen."],
            ["Confirmed / Completed", "Anruf wurde erfolgreich abgeschlossen."],
            ["Escalated", "Notfall- oder Eskalationsprotokoll wurde ausgelöst."],
            ["Cancelled", "Routine oder geplanter Anruf wurde beendet."],
        ],
    },
    {
        "title": "13. Brain Coach",
        "paras": [
            "Brain Coach zeigt kognitive Routinen und Aktivitätsberichte. Die Tabelle zeigt aktive und inaktive Routinen mit Frequenz und bevorzugter Uhrzeit.",
            "Das Report-Icon öffnet den Brain Coach Aktivitätsbericht, nicht das allgemeine Client-Profil. Der Bericht zeigt Durchschnittswert, abgeschlossene Sessions, Fragen, Streak und Details für 7, 30 oder 90 Tage.",
            "Die Profil-Timeline soll die letzte echte Session anzeigen, nicht nur die aktive Einstellung.",
        ],
    },
    {
        "title": "14. Medikation und Adhärenz",
        "paras": [
            "Medication zeigt Medikationssignale und führt zur wöchentlichen Adhärenz pro Client. Medikamente stehen in Zeilen, Wochentage in Spalten.",
            "Der Medikationsdialog umfasst Name, Dosierung, Zweck, Zeiten, Frequenz und ob Erinnerungen aktiv sind.",
            "Vergangene Dosen ohne Eintrag werden als Unconfirmed angezeigt; künftige geplante Dosen bleiben Upcoming.",
        ],
    },
    {
        "title": "15. Risk",
        "paras": [
            "Risk priorisiert operative Arbeit. Die Karten Urgent, Review, No response, Medication und Unassigned müssen erklärbar sein.",
            "Jede Karte sollte eine Hilfe anzeigen, die beschreibt, wie die Zahl berechnet wird. Die Queue aktualisiert sich aus realen Signalen: verpasste Check-ins, unbestätigte Medikation, fehlende Abdeckung, Eskalationen und aktuelle Ereignisse.",
        ],
    },
    {
        "title": "16. Sensors",
        "paras": [
            "Sensors ersetzt die alte Bezeichnung Alerts. Bis der Sensor-Backend-Anschluss fertig ist, zeigt die Seite einen klaren leeren Zustand.",
            "Die Konsole darf keine Sensordaten erfinden. Später können Geräte, Batterie, Verbindung, Alerts und Events ergänzt werden.",
        ],
    },
    {
        "title": "17. VYVA Call Campaigns",
        "paras": [
            "Campaigns ist ein Anruf-first-Arbeitsbereich für VYVA-Kampagnen. Der Ablauf ist: Vorlage wählen, Zielgruppe definieren, Skript schreiben oder generieren, Empfänger prüfen und speichern, planen oder in die Queue stellen.",
            "Vorlagen sind General announcement, Heatwave alert, Vaccination reminder, Scam safety alert, Service update und Create your own.",
            "AI assist kann aus wenigen Worten ein erstes Skript erzeugen. Das Skript bleibt immer editierbar und wird nicht ohne Admin-Bestätigung in die Queue gestellt.",
        ],
    },
    {
        "title": "18. Intelligente Kampagnen-Zielgruppen",
        "paras": [
            "Das Targeting ist regelbasiert. Admins können geografisch, nach Risiko, Gesundheitsbedingung, Versorgungslücke oder zugewiesenem Personal filtern.",
            "Telefonnummer und Zustimmung sind Pflicht-Schutzregeln für VYVA-Anrufkampagnen. Die Vorschau zeigt berechtigte und übersprungene Clients mit Gründen.",
        ],
        "bullets": [
            "Where: Organisation, Land, Stadt oder benannter Bereich.",
            "Who: Risiko-Level und Gesundheitsbedingungen.",
            "Support coverage: alle, unassigned oder einem Provider zugewiesen.",
            "Safeguards: Zustimmung und Telefonnummer bleiben standardmäßig erforderlich.",
        ],
    },
    {
        "title": "19. Team access",
        "paras": [
            "Team access erstellt interne Konsolenkonten für Admins, Koordinatoren und Operatoren. Es erstellt keine Clients.",
            "Superadmin-Berechtigungen werden nicht über diese Seite vergeben. Superadmins bleiben backendverwaltet.",
        ],
    },
    {
        "title": "20. Einstellungen und Sprache",
        "paras": [
            "Settings zeigt aktive Organisation, Standardland, Standardsprache und Zeitzone. Admins können Organisationswerte pflegen, sofern sie die Berechtigung haben.",
            "Die Konsolensprache kann Englisch, Deutsch oder Spanisch sein. Neue Care Profiles übernehmen zunächst die Organisationssprache, können aber angepasst werden.",
        ],
    },
    {
        "title": "21. Reports",
        "paras": [
            "Reports fasst operative Kennzahlen nach Organisation zusammen. Zahlen müssen sich beim Organisationswechsel ändern und dürfen nicht aus einer anderen Organisation gecacht werden.",
            "Reports helfen bei Review, Follow-up, Kampagnenstatus, Brain Coach Verlauf und Versorgungslücken.",
        ],
    },
    {
        "title": "22. Client-Aktivität",
        "paras": [
            "Die Activity timeline zeigt echte Ereignisse: Onboarding, letzte Check-up-Anrufe, letzte Brain Coach Session, Medikationsaktionen, Care coverage Änderungen und Zustimmung.",
            "Einstellungen wie 'Brain Coach active' oder 'Medication plan has items' sind weniger hilfreich als die letzte reale Aktion. Die Timeline soll deshalb Ereignisse statt bloßer Konfigurationen priorisieren.",
        ],
    },
    {
        "title": "23. Datenschutz und minimale Daten",
        "paras": [
            "Die Konsole sollte nur Daten erfassen, die für Betreuung und Koordination erforderlich sind. Vermeiden Sie vollständige medizinische Historien, Versicherungsdaten oder unnötige Diagnosedetails.",
            "Toasts, Logs und Fehlermeldungen dürfen keine Namen, Telefonnummern, Diagnosen, Medikamente oder Freitexte mit sensiblen Angaben enthalten.",
        ],
    },
    {
        "title": "24. Externe Gateways und Integrationen",
        "paras": [
            "Einige Aktionen benötigen externe Gateways: ausgehende Telefonie, WhatsApp, VYVA Voice Connector und zukünftige Sensoren.",
            "Solange ein Gateway nicht verbunden ist, muss die UI das klar sagen und darf keine erfolgreiche Aktion vortäuschen.",
        ],
    },
    {
        "title": "25. Fehlerbehebung",
        "paras": [
            "Wenn der Magic Link nicht ankommt, prüfen Sie Provider-Credits, verifizierte Domain und Absenderkonfiguration.",
            "Wenn eine Seite leer bleibt, sollte die Seite einen Fehlerzustand anzeigen. Prüfen Sie die letzte Veröffentlichung und laden Sie die Seite neu.",
            "Wenn Daten beim Organisationswechsel gleich bleiben, prüfen Sie, ob der Endpoint die aktive Organisation erhält und ob Query-Caches organisationsspezifisch invalidiert werden.",
        ],
    },
]


QUICK_REFERENCE = [
    ["Bedarf", "Wo", "Hinweis"],
    ["Client erstellen", "Clients -> Add client", "Für betreute Personen."],
    ["Clients importieren", "Clients -> Import clients", "CSV mit Organisationsregeln."],
    ["Externes Onboarding prüfen", "Clients -> API intake", "Wenn Daten bereits im Backend existieren."],
    ["Care Profile bearbeiten", "Client-Profil", "Nur berechtigte Rollen."],
    ["Red Cross staff zuweisen", "Profil -> Care coverage", "Nicht mit Emergency contacts mischen."],
    ["Check-ins prüfen", "Check-ins", "Nur Check-up-Anrufe."],
    ["Brain Coach prüfen", "Brain Coach", "Öffnet kognitive Reports."],
    ["Adhärenz anzeigen", "Profil -> View adherence", "Wochenkalender pro Client."],
    ["Kampagne erstellen", "Campaigns -> Create call campaign", "Vorlage, Regeln, Skript und Vorschau."],
    ["Teamkonto erstellen", "Team access", "Nur interne Konten."],
]


SCREENSHOTS = [
    {"section": "2. Anmeldung", "file": "screenshot-01.png", "caption": "Anmeldung per E-Mail und Magic Link."},
    {"section": "2. Anmeldung", "file": "screenshot-02.png", "caption": "Magic-Link-E-Mail mit sicherem Zugang zur Konsole."},
    {"section": "5. Allgemeine Navigation", "file": "screenshot-03.png", "caption": "Seitennavigation mit operativen Bereichen."},
    {"section": "6. Today Dashboard", "file": "screenshot-04.png", "caption": "Today Dashboard mit Kennzahlen, Filtern und Karte."},
    {"section": "7. Clients", "file": "screenshot-05.png", "caption": "Client-Liste mit Suche, Filtern und Care coverage."},
    {"section": "8. Clients erstellen oder importieren", "file": "screenshot-06.png", "caption": "Care Profile Intake für neue Clients."},
    {"section": "9. Client-Profil", "file": "screenshot-07.png", "caption": "Client-Profil mit Status, Aktionen und Care Cards."},
    {"section": "9. Client-Profil", "file": "screenshot-08.png", "caption": "Medikation, Check-ins und Brain Coach im Client-Profil."},
    {"section": "10. Emergency contacts", "file": "screenshot-09.png", "caption": "Emergency contacts aus Onboarding oder Intake."},
    {"section": "11. Red Cross staff", "file": "screenshot-10.png", "caption": "Professionelles Red Cross staff getrennt von persönlichen Kontakten."},
    {"section": "12. Check-ins", "file": "screenshot-11.png", "caption": "Check-in-Routinen mit Status und bevorzugter Uhrzeit."},
    {"section": "13. Brain Coach", "file": "screenshot-12.png", "caption": "Brain Coach Sessions und Aktivitätsberichte."},
    {"section": "14. Medikation und Adhärenz", "file": "screenshot-13.png", "caption": "Wöchentlicher Medikationskalender mit Status pro Dosis."},
    {"section": "15. Risk", "file": "screenshot-14.png", "caption": "Risk Queue mit erklärbaren Kennzahlen."},
    {"section": "17. VYVA Call Campaigns", "file": "screenshot-15.png", "caption": "VYVA Call Campaigns mit Vorlagen und Status."},
    {"section": "18. Intelligente Kampagnen-Zielgruppen", "file": "screenshot-16.png", "caption": "Regelbasierte Zielgruppenauswahl für Kampagnen."},
    {"section": "19. Team access", "file": "screenshot-17.png", "caption": "Team access für interne Konsolenkonten."},
    {"section": "20. Einstellungen und Sprache", "file": "screenshot-18.png", "caption": "Organisation, Sprache und Zeitzone in Settings."},
    {"section": "21. Reports", "file": "screenshot-19.png", "caption": "Reports für operative Kennzahlen."},
    {"section": "22. Client-Aktivität", "file": "screenshot-20.png", "caption": "Activity timeline mit echten Client-Ereignissen."},
]


def ensure_dirs() -> None:
    for directory in [CURRENT_DIR, VERSION_DIR, SOURCE_DIR, PUBLIC_DIR, ARCHIVE_DIR, SCREENSHOT_DIR]:
        directory.mkdir(parents=True, exist_ok=True)


def extract_screenshots() -> None:
    if any(SCREENSHOT_DIR.glob("screenshot-*.png")) or not ENGLISH_MANUAL.exists():
        return
    with ZipFile(ENGLISH_MANUAL) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        for index, name in enumerate(media, start=1):
            extension = Path(name).suffix or ".png"
            (SCREENSHOT_DIR / f"screenshot-{index:02d}{extension}").write_bytes(archive.read(name))


def screenshots_for(section_title: str) -> list[dict[str, str]]:
    return [screenshot for screenshot in SCREENSHOTS if screenshot["section"] == section_title]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "111827") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Segoe UI"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_screenshot(doc: Document, screenshot: dict[str, str]) -> None:
    image_path = SCREENSHOT_DIR / screenshot["file"]
    if not image_path.exists():
        return
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(screenshot["caption"])
    run.italic = True
    run.font.name = "Segoe UI"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    for style_name, size, color in [("Heading 1", 17, "111827"), ("Heading 2", 13, "4F46E5"), ("Heading 3", 11, "111827")]:
        style = styles[style_name]
        style.font.name = "Segoe UI"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ROTES KREUZ")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string("E60012")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VYVA x Red Cross Konsole")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string("111827")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Benutzerhandbuch für Administratoren")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("4F46E5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Version {VERSION} | {UPDATED_DE} | Powered by VYVA")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("6B7280")

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    set_cell_shading(cell, "F5F3FF")
    cell.text = ""
    run = cell.paragraphs[0].add_run(
        "Vollständiges deutsches Handbuch für die VYVA x Red Cross Konsole: Login, Rollen, Organisationen, Clients, Check-ins, Brain Coach, Medikation, Risk, Sensors, Kampagnen, Kontakte, Team, Einstellungen, Reports und Datenschutz."
    )
    run.font.name = "Segoe UI"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("4C1D95")

    doc.add_page_break()
    doc.add_heading("Schnellreferenz", level=1)
    table = doc.add_table(rows=len(QUICK_REFERENCE), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(QUICK_REFERENCE):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=i == 0, color="FFFFFF" if i == 0 else "111827")
            set_cell_shading(cell, "6D4AFF" if i == 0 else ("F9FAFB" if i % 2 == 0 else "FFFFFF"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Handbuch-Link: ").bold = True
    p.add_run(MANUAL_URL)
    doc.add_page_break()

    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)
        for paragraph in section.get("paras", []):
            p = doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(6)
        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
        for index, step in enumerate(section.get("steps", []), start=1):
            doc.add_paragraph(step, style="List Number")
        if "table" in section:
            rows = section["table"]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, row in enumerate(rows):
                for j, value in enumerate(row):
                    cell = table.cell(i, j)
                    set_cell_text(cell, value, bold=i == 0, color="FFFFFF" if i == 0 else "111827")
                    set_cell_shading(cell, "6D4AFF" if i == 0 else ("F9FAFB" if i % 2 == 0 else "FFFFFF"))
        for screenshot in screenshots_for(section["title"]):
            add_docx_screenshot(doc, screenshot)
        doc.add_paragraph()

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Benutzerhandbuch - VYVA x Red Cross Konsole | Powered by VYVA")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("6B7280")

    doc.save(path)


def register_fonts() -> None:
    pdfmetrics.registerFont(TTFont("SegoeUI", "C:/Windows/Fonts/segoeui.ttf"))
    pdfmetrics.registerFont(TTFont("SegoeUI-Bold", "C:/Windows/Fonts/segoeuib.ttf"))


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverTitle", parent=styles["Title"], fontName="SegoeUI-Bold", fontSize=25, leading=31, textColor=colors.HexColor("#111827"), alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="CoverSub", parent=styles["BodyText"], fontName="SegoeUI", fontSize=13, leading=18, textColor=colors.HexColor("#4F46E5"), alignment=TA_CENTER, spaceAfter=6))
    styles.add(ParagraphStyle(name="BodyDE", parent=styles["BodyText"], fontName="SegoeUI", fontSize=10.1, leading=14.8, textColor=colors.HexColor("#374151"), spaceAfter=7))
    styles.add(ParagraphStyle(name="H1DE", parent=styles["Heading1"], fontName="SegoeUI-Bold", fontSize=15, leading=19, textColor=colors.HexColor("#111827"), spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="BulletDE", parent=styles["BodyText"], fontName="SegoeUI", fontSize=9.6, leading=13.6, leftIndent=16, bulletIndent=4, textColor=colors.HexColor("#374151"), spaceAfter=4))
    styles.add(ParagraphStyle(name="NoteDE", parent=styles["BodyText"], fontName="SegoeUI", fontSize=10, leading=15, textColor=colors.HexColor("#4C1D95"), backColor=colors.HexColor("#F5F3FF"), borderColor=colors.HexColor("#DDD6FE"), borderWidth=0.6, borderPadding=8, spaceAfter=10))
    styles.add(ParagraphStyle(name="SmallDE", parent=styles["BodyText"], fontName="SegoeUI", fontSize=8.4, leading=12, textColor=colors.HexColor("#6B7280"), alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="CaptionDE", parent=styles["BodyText"], fontName="SegoeUI", fontSize=8.2, leading=11, textColor=colors.HexColor("#6B7280"), alignment=TA_CENTER, spaceAfter=9))
    return styles


def pdf_screenshot_flowable(image_path: Path) -> PdfImage:
    image = PdfImage(str(image_path))
    max_width = 6.15 * inch
    max_height = 3.75 * inch
    scale = min(max_width / image.imageWidth, max_height / image.imageHeight)
    image.drawWidth = image.imageWidth * scale
    image.drawHeight = image.imageHeight * scale
    image.hAlign = "CENTER"
    return image


def draw_footer(canvas, doc) -> None:
    canvas.saveState()
    width, _ = letter
    canvas.setStrokeColor(colors.HexColor("#E5E7EB"))
    canvas.line(0.75 * inch, 0.55 * inch, width - 0.75 * inch, 0.55 * inch)
    canvas.setFont("SegoeUI", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(0.75 * inch, 0.35 * inch, "Benutzerhandbuch - VYVA x Red Cross Konsole")
    canvas.drawRightString(width - 0.75 * inch, 0.35 * inch, f"Seite {doc.page}")
    canvas.restoreState()


def add_pdf_table(story, rows, styles) -> None:
    data = []
    cols = len(rows[0])
    widths = [1.55 * inch, 2.05 * inch, 2.65 * inch] if cols == 3 else [1.55 * inch, 4.75 * inch]
    for i, row in enumerate(rows):
        data.append([
            Paragraph(
                cell,
                ParagraphStyle(f"tbl{i}{j}{cell[:4]}", fontName="SegoeUI-Bold" if i == 0 else "SegoeUI", fontSize=8.3, leading=11.3, textColor=colors.white if i == 0 else colors.HexColor("#111827")),
            )
            for j, cell in enumerate(row)
        ])
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6D4AFF")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D7DCE8")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(table)


def build_pdf(path: Path) -> None:
    register_fonts()
    styles = pdf_styles()
    story = [
        Spacer(1, 0.28 * inch),
        Paragraph("ROTES KREUZ", ParagraphStyle("RC", fontName="SegoeUI-Bold", fontSize=18, textColor=colors.HexColor("#E60012"), alignment=TA_CENTER, leading=22)),
        Paragraph("VYVA x Red Cross Konsole", styles["CoverTitle"]),
        Paragraph("Benutzerhandbuch für Administratoren", styles["CoverSub"]),
        Paragraph(f"Version {VERSION} | {UPDATED_DE} | Powered by VYVA", styles["SmallDE"]),
        Spacer(1, 0.35 * inch),
        Paragraph("Vollständiges deutsches Handbuch für die aktuelle Konsole: Anmeldung, Rollen, Organisationen, Clients, Check-ins, Brain Coach, Medikation, Risiko, Sensoren, Kampagnen, Kontakte, Team, Einstellungen, Reports und Datenschutz.", styles["NoteDE"]),
        Paragraph(f"Download: {MANUAL_URL}", styles["SmallDE"]),
        PageBreak(),
        Paragraph("Schnellreferenz", styles["H1DE"]),
    ]
    add_pdf_table(story, QUICK_REFERENCE, styles)
    story.append(PageBreak())

    for section in SECTIONS:
        story.append(Paragraph(section["title"], styles["H1DE"]))
        for paragraph in section.get("paras", []):
            story.append(Paragraph(paragraph, styles["BodyDE"]))
        for bullet in section.get("bullets", []):
            story.append(Paragraph(bullet, styles["BulletDE"], bulletText="-"))
        for index, step in enumerate(section.get("steps", []), start=1):
            story.append(Paragraph(f"{index}. {step}", styles["BodyDE"]))
        if "table" in section:
            add_pdf_table(story, section["table"], styles)
        for screenshot in screenshots_for(section["title"]):
            image_path = SCREENSHOT_DIR / screenshot["file"]
            if image_path.exists():
                story.append(Spacer(1, 0.05 * inch))
                story.append(pdf_screenshot_flowable(image_path))
                story.append(Paragraph(screenshot["caption"], styles["CaptionDE"]))
        story.append(Spacer(1, 0.06 * inch))

    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.8 * inch, leftMargin=0.8 * inch, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def write_markdown(path: Path) -> None:
    lines = [
        "# Benutzerhandbuch - VYVA x Red Cross Konsole",
        "",
        f"Version: {VERSION}",
        f"Aktualisiert: {UPDATED_DE}",
        f"URL: {MANUAL_URL}",
        "",
    ]
    for section in SECTIONS:
        lines.append(f"## {section['title']}")
        lines.extend(section.get("paras", []))
        lines.extend([f"- {item}" for item in section.get("bullets", [])])
        lines.extend([f"{i}. {item}" for i, item in enumerate(section.get("steps", []), start=1)])
        if "table" in section:
            rows = section["table"]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        for screenshot in screenshots_for(section["title"]):
            lines.append("")
            lines.append(f"![{screenshot['caption']}](../current/screenshots-from-english/{screenshot['file']})")
            lines.append(f"*{screenshot['caption']}*")
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def update_manifest() -> None:
    metadata_path = PUBLIC_DIR / "manual-version.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    languages = metadata.setdefault("languages", {})
    languages["de"] = {
        "title": "Benutzerhandbuch - VYVA x Red Cross Konsole",
        "latestPdf": "/manuals/VYVA_Admin_Console_User_Manual_DE.pdf",
        "archivePdf": "/manuals/archive/VYVA_Admin_Console_User_Manual_DE_2026-06-20.pdf",
    }
    metadata["latestGermanPdf"] = "/manuals/VYVA_Admin_Console_User_Manual_DE.pdf"
    metadata["germanManualUrl"] = MANUAL_URL
    metadata_path.write_text(json.dumps(metadata, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def main() -> None:
    ensure_dirs()
    extract_screenshots()
    docx_root = ROOT / f"{OUT_NAME}.docx"
    pdf_root = ROOT / f"{OUT_NAME}.pdf"
    build_docx(docx_root)
    build_pdf(pdf_root)

    for target in [CURRENT_DIR / f"{OUT_NAME}.docx", VERSION_DIR / f"{OUT_NAME}.docx"]:
        shutil.copy2(docx_root, target)

    for target in [
        CURRENT_DIR / f"{OUT_NAME}.pdf",
        VERSION_DIR / f"{OUT_NAME}.pdf",
        PUBLIC_DIR / f"{OUT_NAME}.pdf",
        ARCHIVE_DIR / f"{OUT_NAME}_{VERSION}.pdf",
    ]:
        shutil.copy2(pdf_root, target)

    write_markdown(SOURCE_DIR / "manual-german-2026-06-20.md")
    update_manifest()
    print(f"created {docx_root}")
    print(f"created {pdf_root}")


if __name__ == "__main__":
    main()
