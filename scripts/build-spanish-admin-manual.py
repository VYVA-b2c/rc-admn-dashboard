from __future__ import annotations

import json
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path.cwd()
VERSION = "2026-06-20"
UPDATED_ES = "20 de junio de 2026"
MANUAL_URL = "https://rcadmin.vyva.life/manuals/VYVA_Admin_Console_User_Manual_ES.pdf"

OUT_NAME = "VYVA_Admin_Console_User_Manual_ES"
CURRENT_DIR = ROOT / "docs" / "admin-manual" / "current"
VERSION_DIR = ROOT / "docs" / "admin-manual" / "versions" / VERSION
SOURCE_DIR = ROOT / "docs" / "admin-manual" / "source"
PUBLIC_DIR = ROOT / "public" / "manuals"
ARCHIVE_DIR = PUBLIC_DIR / "archive"
SCREENSHOT_DIR = CURRENT_DIR / "screenshots-from-english"
ENGLISH_MANUAL = ROOT / "VYVA_Admin_Console_User_Manual.docx"


def ensure_dirs() -> None:
    for directory in [CURRENT_DIR, VERSION_DIR, SOURCE_DIR, PUBLIC_DIR, ARCHIVE_DIR, SCREENSHOT_DIR]:
        directory.mkdir(parents=True, exist_ok=True)


SECTIONS = [
    {
        "title": "1. Propósito y alcance de la consola",
        "paras": [
            "La consola VYVA x Cruz Roja es un espacio operativo para coordinar seguimiento, llamadas, señales de riesgo, medicación, campañas y cobertura de cuidado para clientes mayores atendidos por una organización de Cruz Roja.",
            "Este manual está escrito para administradores, coordinadores, operadores y personal profesional autorizado. Explica cómo usar la consola actual, qué datos debe mostrar cada área y qué acciones requieren integración externa antes de ejecutarse.",
            "La consola debe usarse con criterio de mínima información necesaria: registrar únicamente los datos que ayudan a coordinar cuidado, seguimiento y seguridad. No debe convertirse en una historia clínica completa ni en un repositorio de información médica no operativa.",
        ],
        "bullets": [
            "Clientes son las personas atendidas por Cruz Roja.",
            "Miembros del equipo son usuarios internos con acceso a la consola.",
            "Contactos de emergencia son familiares, vecinos, cuidadores o contactos personales capturados durante onboarding o intake.",
            "Personal Cruz Roja es personal profesional asignable a clientes.",
        ],
    },
    {
        "title": "2. Inicio de sesión",
        "paras": [
            "El acceso visible es solo por correo electrónico. El administrador introduce su email y recibe un enlace mágico de inicio de sesión. Google y Microsoft están ocultos por ahora para evitar confusión hasta que se activen correctamente.",
            "Supabase Auth se usa como capa de identidad para el enlace mágico. Replit mantiene los datos operativos: organizaciones, roles, equipos, clientes, campañas y asignaciones.",
            "Si el email no tiene acceso concedido, la consola debe bloquear el acceso después del intento de inicio de sesión con un mensaje claro. El sistema no debe revelar detalles sensibles antes de enviar el enlace.",
        ],
        "steps": [
            "Abrir la URL publicada de la consola.",
            "Introducir el email de administrador.",
            "Revisar el correo recibido desde el dominio VYVA configurado.",
            "Abrir el enlace mágico más reciente.",
            "Confirmar que la organización mostrada es la correcta.",
        ],
    },
    {
        "title": "3. Roles y permisos",
        "paras": [
            "Los permisos se resuelven a partir del email autenticado y de los registros del backend. La interfaz debe mostrar el rol real del usuario, por ejemplo Super admin, Admin, Coordinador u Operador. No debe mostrar Operator si el usuario es Super admin.",
            "Los superadmins no se crean desde la interfaz. Se conceden desde controles backend, como variables de entorno o campos de perfil administrados. Esto reduce el riesgo de elevar permisos por error.",
        ],
        "table": [
            ["Rol", "Alcance principal", "Notas"],
            ["Super admin", "Puede cambiar de organización y crear administradores por organización.", "No se gestiona desde la UI."],
            ["Admin", "Gestiona clientes, equipo, campañas, contactos, horarios y configuración de su organización.", "Puede crear miembros de equipo."],
            ["Coordinador / operador", "Revisa colas, clientes y tareas asignadas.", "Las acciones sensibles pueden ser de solo lectura."],
            ["Personal Cruz Roja asignado", "Puede editar planes o rutinas del cliente cuando es responsable principal y hay consentimiento.", "No equivale a contacto de emergencia."],
        ],
    },
    {
        "title": "4. Organización activa",
        "paras": [
            "La organización activa controla qué datos se muestran en toda la consola. Al cambiar de Red Cross Zamora a Red Cross Leipzig, deben cambiar los clientes, campañas, check-ins, Brain Coach, medicación, riesgo, contactos, personal, informes y mapa.",
            "Red Cross Zamora usa español, España y zona horaria Europe/Madrid por defecto. Red Cross Leipzig usa alemán, Alemania y Europe/Berlin por defecto. El usuario puede cambiar idioma de interfaz si tiene permiso.",
            "La asignación automática debe ser estricta: números y direcciones españolas pertenecen a Zamora; números y direcciones alemanas pertenecen a Leipzig. Un cliente no debe aparecer en otra organización por fallback genérico.",
        ],
        "bullets": [
            "El selector de organización aparece para superadmins o usuarios con más de una organización.",
            "El mapa debe centrarse en el país de la organización activa.",
            "Los endpoints deben recibir o resolver el contexto de organización antes de devolver datos.",
        ],
    },
    {
        "title": "5. Navegación general",
        "paras": [
            "La barra lateral agrupa las áreas en Principal, Seguimiento y Gestión. Las etiquetas deben ser consistentes: Clientes para personas atendidas, Riesgo para priorización operativa, Sensores para datos de dispositivos, Check-ins para llamadas de seguimiento, Brain Coach para sesiones cognitivas y Campañas para llamadas VYVA.",
            "El encabezado superior muestra el nombre de la consola, la organización activa, el estado del sistema y el usuario conectado. Si hay selector de organización, el cambio debe refrescar los datos de la vista actual.",
        ],
        "table": [
            ["Área", "Uso"],
            ["Today", "Vista operativa diaria con métricas y mapa."],
            ["Clients", "Cola de clientes y perfiles de cuidado."],
            ["Risk", "Casos priorizados por riesgo, revisión, medicación, no respuesta o falta de asignación."],
            ["Sensors", "Espacio reservado para datos de dispositivos cuando el backend esté listo."],
            ["Check-ins", "Llamadas de seguimiento y su estado."],
            ["Brain Coach", "Rutinas y reportes de actividad cognitiva."],
            ["Medication", "Seguimiento de medicación y adherencia."],
            ["Campaigns", "Campañas de llamadas VYVA."],
        ],
    },
    {
        "title": "6. Panel Today",
        "paras": [
            "Today resume el estado operativo de la organización activa. Las tarjetas superiores deben mostrar números calculados con lógica clara y periodo definido. Cuando el usuario hace clic en una tarjeta, debe poder ver la lista que explica el número cuando aplique.",
            "El mapa debe estar arriba y debe mantener el comportamiento Leaflet: tamaño estable, resize correcto, tiles, clustering y viewport coherente. Debe mostrar clientes, oficinas y equipo de campo como capas.",
            "Las métricas de check-ins deben ser semanales: completados frente a esperados durante la semana. No basta con mostrar 1/1 si la rutina es diaria y se esperan siete llamadas semanales.",
        ],
        "bullets": [
            "Inscritos: clientes activos en la organización.",
            "Urgente: clientes con señales que requieren acción inmediata.",
            "Revisar: clientes que necesitan evaluación del operador hoy.",
            "Check-ins: cumplimiento semanal real.",
            "Medicación: dosis o confirmaciones pendientes.",
            "Sin asignar: clientes sin cobertura operativa suficiente.",
        ],
    },
    {
        "title": "7. Clientes",
        "paras": [
            "Clientes es la cola principal de personas atendidas. Desde aquí se puede buscar por nombre, teléfono, ciudad o cuidador, filtrar por urgencia, revisión, sin respuesta, medicación, check-ins o sin asignar, y abrir el perfil del cliente.",
            "La página tiene tres entradas de creación: Add client para crear uno a uno, Import clients para CSV y API intake para usuarios que llegan desde onboarding externo. Estas acciones son para clientes, no para miembros del equipo.",
        ],
        "steps": [
            "Usar Add client cuando un admin crea un cliente directamente.",
            "Usar Import clients cuando se cargan varios clientes con CSV.",
            "Usar API intake cuando los datos ya existen en el backend de onboarding.",
            "Abrir una fila para revisar el perfil y la cobertura de cuidado.",
        ],
    },
    {
        "title": "8. Crear o importar clientes",
        "paras": [
            "El formulario de cliente debe recoger datos operativos de cuidado en secciones claras: persona y contacto, perfil médico mínimo, medicación, consentimiento, contacto de emergencia y rutinas de seguimiento.",
            "El teléfono debe mostrar formato internacional claro, empezando por + y código de país. Esto ayuda a enrutar correctamente por organización y a evitar números inválidos.",
            "El CSV debe respetar las mismas reglas de mínima información necesaria y nunca crear miembros de equipo por error. Los contactos de emergencia importados deben quedar como contactos personales, no como personal Cruz Roja.",
        ],
        "bullets": [
            "Nombre y apellidos son obligatorios.",
            "Teléfono debe estar en formato internacional.",
            "Consentimiento controla llamadas rutinarias.",
            "Solo se pide información médica útil para coordinación.",
            "Sensores quedan fuera hasta que el backend esté listo.",
        ],
    },
    {
        "title": "9. Perfil del cliente",
        "paras": [
            "El perfil del cliente es la vista de trabajo principal para una persona atendida. Debe mostrar nombre, ciudad, estado, teléfono, idioma, dirección, consentimiento, último contacto y cobertura de cuidado sin textos de relleno como Unknown cuando exista información real.",
            "La tarjeta de medicación y check-ins debe mostrar medicamentos, horarios, frecuencia y botones de edición si el usuario tiene permisos. Ver adherencia abre el calendario semanal de medicación por cliente.",
            "Las acciones Call now, Send WhatsApp y Contact care provider deben informar que se requiere conexión de pasarela cuando todavía no hay integración activa. No deben simular una llamada o mensaje realizado.",
        ],
        "bullets": [
            "Datos clave: edad, teléfono, idioma, dirección, consentimiento y último contacto.",
            "Salud y cuidado: condiciones, movilidad y notas mínimas de seguridad.",
            "Medicación y check-ins: medicamentos, horarios y rutinas.",
            "Cobertura: contacto de emergencia y personal Cruz Roja principal.",
            "Actividad: eventos reales ordenados por fecha.",
        ],
    },
    {
        "title": "10. Contactos de emergencia",
        "paras": [
            "Los contactos de emergencia son personas de apoyo no profesionales: familiares, vecinos, cuidadores informales u otros contactos personales. Se capturan durante onboarding, llamadas entrantes o el formulario de cliente.",
            "La página Contactos de emergencia debe listar solo estos contactos personales. No debe mezclar personal Cruz Roja. La tabla debe incluir nombre, relación o rol, teléfono, fuente, clientes vinculados y número de asignaciones.",
            "Un cliente puede tener varios contactos de emergencia y un contacto puede estar vinculado a varios clientes. La vista de perfil debe evitar repetir el mismo contacto como resumen y como detalle de forma confusa.",
        ],
    },
    {
        "title": "11. Personal Cruz Roja",
        "paras": [
            "El personal Cruz Roja representa profesionales o equipo operativo. Debe gestionarse separado de los contactos de emergencia. La asignación de personal debe pedir un rol profesional mediante desplegable, no una relación familiar.",
            "Un cliente puede tener personal profesional asignado. El sistema debe diferenciar el personal principal de otros apoyos profesionales. Esta asignación puede habilitar permisos de edición cuando el usuario asignado es responsable principal.",
        ],
        "bullets": [
            "Ejemplos de rol: coordinador de campo, operador principal, apoyo de medicación, trabajador social, supervisor.",
            "No usar Daughter, neighbor o caregiver como rol de personal profesional.",
            "Si no hay personal disponible, la UI debe indicar que falta crear registros de personal para la organización activa.",
        ],
    },
    {
        "title": "12. Check-ins",
        "paras": [
            "Check-ins es solo para llamadas de seguimiento o check-up calls. Brain Coach no debe aparecer en esta página. Cada fila debe mostrar cliente, teléfono, tipo, estado, último check-in, frecuencia, hora preferida y acciones.",
            "El último check-in debe venir de actividad real. Si la llamada programada ya pasó y no hay registro de éxito, debe mostrar Missed today o equivalente. Si fue confirmada, debe mostrar Confirmed o Completed. Si se activó protocolo de emergencia, debe mostrarse Escalated.",
            "Los admins y el personal Cruz Roja principal pueden editar horarios cuando hay consentimiento. Los contactos de emergencia no pueden editar rutinas.",
        ],
        "table": [
            ["Estado", "Significado"],
            ["Active", "La rutina está habilitada."],
            ["Missed", "La llamada esperada pasó sin confirmación."],
            ["Confirmed / Completed", "La llamada fue completada con resultado válido."],
            ["Escalated", "Se activó protocolo de emergencia o revisión crítica."],
            ["Cancelled", "La rutina o llamada fue cancelada correctamente."],
        ],
    },
    {
        "title": "13. Brain Coach",
        "paras": [
            "Brain Coach tiene una página separada de sesiones. La tabla muestra clientes con rutina cognitiva activa o inactiva, teléfono, frecuencia, hora preferida y acciones.",
            "El icono de reporte debe abrir la página de reporte de actividad cognitiva, no el perfil general. El reporte muestra promedio, sesiones completadas, total de preguntas, racha y detalle por periodo de 7, 30 o 90 días.",
            "La actividad del perfil debe mostrar la última sesión registrada, no solo que la rutina está activa. Si no hay sesiones, debe mostrar un estado claro como sin historial todavía.",
        ],
    },
    {
        "title": "14. Medicación y adherencia",
        "paras": [
            "Medicación permite revisar señales relacionadas con medicamentos y abrir la adherencia semanal por cliente. La página de adherencia muestra medicamentos por fila y días por columna.",
            "El formulario de medicación debe incluir nombre, dosis, propósito, horarios, frecuencia y si los recordatorios están activos. La frecuencia es necesaria para interpretar correctamente recordatorios y adherencia.",
            "Los estados del calendario son tomado, perdido, sin confirmar y próximo. Las dosis pasadas sin registro deben aparecer como sin confirmar, no como próximas.",
        ],
        "bullets": [
            "Admins pueden editar medicamentos.",
            "Personal Cruz Roja principal puede editar si tiene permiso operativo.",
            "Los contactos de emergencia no editan medicación.",
            "No incluir información médica extensa si no es necesaria para coordinación.",
        ],
    },
    {
        "title": "15. Riesgo",
        "paras": [
            "Riesgo prioriza trabajo operativo. La página debe mostrar tarjetas de Urgente, Revisión, Sin respuesta, Medicación y Sin asignar. Cada tarjeta debe incluir una ayuda con explicación de cómo se calcula el número.",
            "La cola debe actualizarse según señales reales: check-ins perdidos, medicación sin confirmar, falta de cobertura, escalaciones y eventos recientes. Si no hay casos, debe explicar qué señales generarían entradas.",
        ],
        "bullets": [
            "Urgente: requiere acción inmediata.",
            "Revisión: requiere evaluación del operador.",
            "Sin respuesta: no se obtuvo confirmación en rutinas o campañas.",
            "Medicación: señales de adherencia o confirmación pendiente.",
            "Sin asignar: falta responsable o cobertura.",
        ],
    },
    {
        "title": "16. Sensores",
        "paras": [
            "Sensores reemplaza la etiqueta Alerts. Por ahora puede existir como sección preparada para dispositivos y señales futuras. Si el backend de sensores no existe, la UI debe explicar claramente que todavía no hay datos conectados.",
            "No debe inventar datos de sensores. Cuando haya integración, la página podrá mostrar dispositivos, alertas, batería, conexión y eventos recientes.",
        ],
    },
    {
        "title": "17. Campañas de llamadas VYVA",
        "paras": [
            "Campañas es un espacio de llamadas VYVA, no una herramienta genérica multicanal. El flujo correcto es seleccionar una plantilla, definir audiencia, escribir o generar el guion, revisar destinatarios y guardar, programar o poner en cola.",
            "Las plantillas disponibles incluyen anuncio general, alerta de ola de calor, recordatorio de vacunación, alerta de estafas, actualización de servicios y crear tu propia campaña.",
            "La opción de IA ayuda al admin a escribir unas palabras sobre el objetivo y generar un primer borrador del guion. El texto siempre debe ser editable antes de guardar o poner en cola.",
        ],
        "steps": [
            "Elegir plantilla o Crear tu propia.",
            "Definir reglas de público objetivo.",
            "Editar el guion de llamada.",
            "Previsualizar destinatarios elegibles y omitidos.",
            "Guardar borrador, programar llamadas o poner en cola.",
        ],
    },
    {
        "title": "18. Segmentación inteligente de campañas",
        "paras": [
            "La segmentación debe ser práctica y entendible para administradores. En lugar de un texto libre llamado target summary, la consola debe ofrecer reglas claras: geografía, riesgo, condición de salud, proveedor asignado, consentimiento y teléfono.",
            "Consentimiento y teléfono deben estar activados por defecto como salvaguardas. La previsualización debe mostrar cuántos clientes son elegibles y cuántos se omiten, con razones como sin teléfono, sin consentimiento, fuera del área, sin condición seleccionada o proveedor no coincidente.",
        ],
        "bullets": [
            "Dónde: toda la organización, país, ciudad o área.",
            "Quién: todos, estable, revisión, urgente o alto riesgo.",
            "Salud: condiciones registradas en el perfil del cliente.",
            "Cobertura: todos, sin asignar o asignados a un proveedor concreto.",
            "Salvaguardas: teléfono y consentimiento.",
        ],
    },
    {
        "title": "19. Team access",
        "paras": [
            "Team access sirve para crear cuentas de personal con acceso a la consola. No crea clientes. El texto debe decir Add team member y explicar que es para staff console access only.",
            "Los admins de organización pueden crear miembros del equipo dentro de su organización. Los superadmins pueden crear administradores por organización. La creación o promoción de superadmins queda fuera de la UI.",
        ],
    },
    {
        "title": "20. Configuración e idioma",
        "paras": [
            "Configuración muestra la organización activa, país, idioma por defecto y zona horaria. Los admins pueden cambiar valores permitidos según rol. El idioma de la consola puede ser inglés, alemán o español.",
            "La página debe tener una forma clara de volver a la consola. Los cambios de organización o idioma deben actualizar la UI sin dejar datos mezclados de otra organización.",
        ],
    },
    {
        "title": "21. Informes",
        "paras": [
            "Informes resume actividad útil para coordinación. Puede incluir campañas, check-ins, medicación, servicios, estado de población y tendencias por organización.",
            "Todos los informes deben respetar la organización activa. Si el superadmin cambia de Zamora a Leipzig, los datos deben cambiar también.",
        ],
    },
    {
        "title": "22. Actividad del cliente",
        "paras": [
            "La actividad del cliente debe registrar eventos reales, no solo configuraciones activas. Por ejemplo, debe mostrar último check-in realizado o perdido, última sesión Brain Coach registrada, último evento de medicación, consentimiento registrado, asignación de cuidado y cambios relevantes.",
            "No es útil mostrar únicamente 'Brain Coach sessions active' o 'Medication plan has 2 items' como actividad. Eso describe configuración, no actividad. La línea temporal debe priorizar acciones y eventos.",
        ],
    },
    {
        "title": "23. Comunicaciones y pasarelas",
        "paras": [
            "Hasta que existan conectores de voz, WhatsApp o proveedor de llamadas, los botones de comunicación deben decir que se requiere una conexión de pasarela. Esto evita que el operador crea que se llamó, se escribió o se contactó a alguien cuando no ocurrió.",
            "Las campañas pueden preparar y poner en cola trabajos, pero no deben iniciar llamadas reales si el conector de voz no está configurado.",
        ],
    },
    {
        "title": "24. Seguridad, privacidad y mínimo necesario",
        "paras": [
            "La consola debe seguir una regla de mínimo necesario. Registrar solo lo que ayuda al equipo a coordinar cuidado y seguimiento. Evitar diagnósticos narrativos extensos, datos financieros, aseguradoras o información clínica no operativa.",
            "Los mensajes de éxito o error no deben incluir nombres, teléfonos, medicaciones o condiciones médicas. Los logs tampoco deben exponer PHI innecesaria.",
        ],
        "bullets": [
            "No guardar borradores médicos en localStorage o sessionStorage.",
            "No registrar datos sensibles en toasts o consola.",
            "Validar consentimiento antes de llamadas rutinarias.",
            "Separar contactos personales de personal profesional.",
            "Mantener superadmin fuera de flujos de UI normales.",
        ],
    },
    {
        "title": "25. Solución de problemas",
        "paras": [
            "Si no llega el enlace mágico, revisar que el proveedor de correo tenga créditos, dominio verificado y remitente configurado. La consola debe mostrar errores concretos como falta de remitente, límite de créditos o espera temporal.",
            "Si una página queda en blanco, revisar el error de renderizado, refrescar la página y confirmar que la última publicación se completó correctamente. Las páginas críticas deben mostrar un estado de error claro en lugar de quedar vacías.",
            "Si los datos no cambian al cambiar de organización, revisar que el endpoint recibe el contexto de organización y que no existe caché compartida entre organizaciones.",
        ],
    },
]


QUICK_REFERENCE = [
    ["Necesidad", "Dónde hacerlo", "Nota"],
    ["Crear cliente", "Clients -> Add client", "Para personas atendidas."],
    ["Importar clientes", "Clients -> Import clients", "Carga CSV con reglas de organización."],
    ["Revisar onboarding externo", "Clients -> API intake", "Cuando el backend ya contiene datos."],
    ["Editar perfil de cuidado", "Perfil del cliente", "Solo roles autorizados."],
    ["Asignar personal Cruz Roja", "Perfil -> Care coverage", "No mezclar con contactos de emergencia."],
    ["Ver check-ins", "Check-ins", "Solo llamadas de seguimiento."],
    ["Ver Brain Coach", "Brain Coach", "Abre reportes cognitivos."],
    ["Ver adherencia", "Perfil -> View adherence", "Calendario semanal por cliente."],
    ["Crear campaña", "Campaigns -> Create call campaign", "Plantilla, reglas, guion y preview."],
    ["Crear equipo", "Team access", "Solo cuentas internas."],
]


SCREENSHOTS = [
    {"section": "2. Inicio de sesión", "file": "screenshot-01.png", "caption": "Pantalla de inicio de sesión con acceso por enlace mágico."},
    {"section": "2. Inicio de sesión", "file": "screenshot-02.png", "caption": "Correo de acceso seguro con botón para abrir la consola VYVA."},
    {"section": "5. Navegación general", "file": "screenshot-03.png", "caption": "Barra lateral con las áreas principales de la consola."},
    {"section": "6. Panel Today", "file": "screenshot-04.png", "caption": "Panel Today con métricas operativas, filtros y mapa superior."},
    {"section": "7. Clientes", "file": "screenshot-05.png", "caption": "Lista de clientes con búsqueda, filtros y cobertura de cuidado."},
    {"section": "8. Crear o importar clientes", "file": "screenshot-06.png", "caption": "Formulario de alta de cliente con datos de contacto y perfil de cuidado."},
    {"section": "9. Perfil del cliente", "file": "screenshot-07.png", "caption": "Perfil del cliente con datos clave, estado, acciones y tarjetas de cuidado."},
    {"section": "9. Perfil del cliente", "file": "screenshot-08.png", "caption": "Tarjeta de medicación, check-ins y Brain Coach del cliente."},
    {"section": "10. Contactos de emergencia", "file": "screenshot-09.png", "caption": "Directorio de contactos de emergencia capturados durante onboarding o intake."},
    {"section": "11. Personal Cruz Roja", "file": "screenshot-10.png", "caption": "Asignación de personal profesional Cruz Roja separada de contactos personales."},
    {"section": "12. Check-ins", "file": "screenshot-11.png", "caption": "Vista de check-ins con frecuencia, hora preferida y último estado registrado."},
    {"section": "13. Brain Coach", "file": "screenshot-12.png", "caption": "Sesiones Brain Coach y acceso al reporte de actividad cognitiva."},
    {"section": "14. Medicación y adherencia", "file": "screenshot-13.png", "caption": "Calendario semanal de adherencia por medicamento y estado de dosis."},
    {"section": "15. Riesgo", "file": "screenshot-14.png", "caption": "Cola de riesgo con tarjetas explicables y filtros operativos."},
    {"section": "17. Campañas de llamadas VYVA", "file": "screenshot-15.png", "caption": "Campañas de llamadas VYVA con plantillas, estados y acciones de cola."},
    {"section": "18. Segmentación inteligente de campañas", "file": "screenshot-16.png", "caption": "Constructor de audiencia por geografía, riesgo, condiciones y cobertura."},
    {"section": "19. Team access", "file": "screenshot-17.png", "caption": "Gestión de acceso para miembros del equipo, separada de clientes."},
    {"section": "20. Configuración e idioma", "file": "screenshot-18.png", "caption": "Configuración de organización activa, idioma y zona horaria."},
    {"section": "21. Informes", "file": "screenshot-19.png", "caption": "Informes operativos para revisar actividad y métricas por organización."},
    {"section": "22. Actividad del cliente", "file": "screenshot-20.png", "caption": "Línea de tiempo con acciones reales: onboarding, llamadas, medicación y sesiones."},
]


def screenshots_for(section_title: str) -> list[dict[str, str]]:
    return [screenshot for screenshot in SCREENSHOTS if screenshot["section"] == section_title]


def extract_screenshots() -> None:
    if any(SCREENSHOT_DIR.glob("screenshot-*.png")):
        return
    if not ENGLISH_MANUAL.exists():
        return
    with ZipFile(ENGLISH_MANUAL) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        for index, name in enumerate(media, start=1):
            extension = Path(name).suffix or ".png"
            (SCREENSHOT_DIR / f"screenshot-{index:02d}{extension}").write_bytes(archive.read(name))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "111827") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Segoe UI"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_screenshot(doc: Document, screenshot: dict[str, str]) -> None:
    image_path = SCREENSHOT_DIR / screenshot["file"]
    if not image_path.exists():
        return
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(screenshot["caption"])
    run.italic = True
    run.font.name = "Segoe UI"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    for style_name, size, color in [
        ("Heading 1", 17, "111827"),
        ("Heading 2", 13, "4F46E5"),
        ("Heading 3", 11, "111827"),
    ]:
        style = styles[style_name]
        style.font.name = "Segoe UI"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CRUZ ROJA")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string("E60012")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Consola VYVA x Cruz Roja")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string("111827")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de usuario para administradores")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("4F46E5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Versión {VERSION} | {UPDATED_ES} | Powered by VYVA")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("6B7280")

    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_shading(cell, "F5F3FF")
    cell.text = ""
    r = cell.paragraphs[0].add_run(
        "Manual completo en español para operar la consola: acceso, roles, organizaciones, clientes, check-ins, Brain Coach, medicación, riesgo, sensores, campañas, contactos, equipo, configuración, informes y privacidad."
    )
    r.font.name = "Segoe UI"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("4C1D95")

    doc.add_page_break()
    doc.add_heading("Referencia rápida", level=1)
    table = doc.add_table(rows=len(QUICK_REFERENCE), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(QUICK_REFERENCE):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=i == 0, color="FFFFFF" if i == 0 else "111827")
            set_cell_shading(cell, "6D4AFF" if i == 0 else ("F9FAFB" if i % 2 == 0 else "FFFFFF"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Enlace del manual: ").bold = True
    p.add_run(MANUAL_URL)
    doc.add_page_break()

    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)
        for paragraph in section.get("paras", []):
            p = doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(6)
        for bullet in section.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(bullet)
        for index, step in enumerate(section.get("steps", []), start=1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(step)
        if "table" in section:
            rows = section["table"]
            cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, row in enumerate(rows):
                for j, value in enumerate(row):
                    cell = table.cell(i, j)
                    set_cell_text(cell, value, bold=i == 0, color="FFFFFF" if i == 0 else "111827")
                    set_cell_shading(cell, "6D4AFF" if i == 0 else ("F9FAFB" if i % 2 == 0 else "FFFFFF"))
        for screenshot in screenshots_for(section["title"]):
            add_docx_screenshot(doc, screenshot)
        doc.add_paragraph()

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Manual de usuario - Consola VYVA x Cruz Roja | Powered by VYVA")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("6B7280")

    doc.save(path)


def register_fonts() -> None:
    pdfmetrics.registerFont(TTFont("SegoeUI", "C:/Windows/Fonts/segoeui.ttf"))
    pdfmetrics.registerFont(TTFont("SegoeUI-Bold", "C:/Windows/Fonts/segoeuib.ttf"))


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontName="SegoeUI-Bold",
            fontSize=26,
            leading=32,
            textColor=colors.HexColor("#111827"),
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSub",
            parent=styles["Normal"],
            fontName="SegoeUI",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#4F46E5"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyES",
            parent=styles["BodyText"],
            fontName="SegoeUI",
            fontSize=10.1,
            leading=14.8,
            textColor=colors.HexColor("#374151"),
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1ES",
            parent=styles["Heading1"],
            fontName="SegoeUI-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#111827"),
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletES",
            parent=styles["BodyText"],
            fontName="SegoeUI",
            fontSize=9.6,
            leading=13.6,
            leftIndent=16,
            bulletIndent=4,
            textColor=colors.HexColor("#374151"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="NoteES",
            parent=styles["BodyText"],
            fontName="SegoeUI",
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#4C1D95"),
            backColor=colors.HexColor("#F5F3FF"),
            borderColor=colors.HexColor("#DDD6FE"),
            borderWidth=0.6,
            borderPadding=8,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallES",
            parent=styles["BodyText"],
            fontName="SegoeUI",
            fontSize=8.4,
            leading=12,
            textColor=colors.HexColor("#6B7280"),
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CaptionES",
            parent=styles["BodyText"],
            fontName="SegoeUI",
            fontSize=8.2,
            leading=11,
            textColor=colors.HexColor("#6B7280"),
            alignment=TA_CENTER,
            spaceAfter=9,
        )
    )
    return styles


def pdf_screenshot_flowable(image_path: Path) -> PdfImage:
    image = PdfImage(str(image_path))
    max_width = 6.15 * inch
    max_height = 3.75 * inch
    scale = min(max_width / image.imageWidth, max_height / image.imageHeight)
    image.drawWidth = image.imageWidth * scale
    image.drawHeight = image.imageHeight * scale
    image.hAlign = "CENTER"
    return image


def draw_footer(canvas, doc) -> None:
    canvas.saveState()
    width, _ = letter
    canvas.setStrokeColor(colors.HexColor("#E5E7EB"))
    canvas.line(0.75 * inch, 0.55 * inch, width - 0.75 * inch, 0.55 * inch)
    canvas.setFont("SegoeUI", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(0.75 * inch, 0.35 * inch, "Manual de usuario - Consola VYVA x Cruz Roja")
    canvas.drawRightString(width - 0.75 * inch, 0.35 * inch, f"Página {doc.page}")
    canvas.restoreState()


def build_pdf(path: Path) -> None:
    register_fonts()
    styles = pdf_styles()
    story = []
    story.append(Spacer(1, 0.28 * inch))
    story.append(Paragraph("CRUZ ROJA", ParagraphStyle("RC", fontName="SegoeUI-Bold", fontSize=18, textColor=colors.HexColor("#E60012"), alignment=TA_CENTER, leading=22)))
    story.append(Paragraph("Consola VYVA x Cruz Roja", styles["CoverTitle"]))
    story.append(Paragraph("Manual de usuario para administradores", styles["CoverSub"]))
    story.append(Paragraph(f"Versión {VERSION} | {UPDATED_ES} | Powered by VYVA", styles["SmallES"]))
    story.append(Spacer(1, 0.35 * inch))
    story.append(
        Paragraph(
            "Manual completo en español para operar la consola: acceso, roles, organizaciones, clientes, check-ins, Brain Coach, medicación, riesgo, sensores, campañas, contactos, equipo, configuración, informes y privacidad.",
            styles["NoteES"],
        )
    )
    story.append(Paragraph(f"Descarga: {MANUAL_URL}", styles["SmallES"]))
    story.append(PageBreak())

    story.append(Paragraph("Referencia rápida", styles["H1ES"]))
    data = []
    for i, row in enumerate(QUICK_REFERENCE):
        data.append(
            [
                Paragraph(
                    cell,
                    ParagraphStyle(
                        f"qr{i}{j}",
                        fontName="SegoeUI-Bold" if i == 0 else "SegoeUI",
                        fontSize=8.4,
                        leading=11.5,
                        textColor=colors.white if i == 0 else colors.HexColor("#111827"),
                    ),
                )
                for j, cell in enumerate(row)
            ]
        )
    table = Table(data, colWidths=[1.55 * inch, 2.05 * inch, 2.65 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6D4AFF")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D7DCE8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
            ]
        )
    )
    story.append(table)
    story.append(PageBreak())

    for section in SECTIONS:
        story.append(Paragraph(section["title"], styles["H1ES"]))
        for paragraph in section.get("paras", []):
            story.append(Paragraph(paragraph, styles["BodyES"]))
        for bullet in section.get("bullets", []):
            story.append(Paragraph(bullet, styles["BulletES"], bulletText="-"))
        for index, step in enumerate(section.get("steps", []), start=1):
            story.append(Paragraph(f"{index}. {step}", styles["BodyES"]))
        if "table" in section:
            table_data = []
            column_count = len(section["table"][0])
            widths = [1.4 * inch, 2.15 * inch, 2.75 * inch] if column_count == 3 else [1.55 * inch, 4.75 * inch]
            for i, row in enumerate(section["table"]):
                table_data.append(
                    [
                        Paragraph(
                            cell,
                            ParagraphStyle(
                                f"t{section['title']}{i}{j}",
                                fontName="SegoeUI-Bold" if i == 0 else "SegoeUI",
                                fontSize=8.3,
                                leading=11.3,
                                textColor=colors.white if i == 0 else colors.HexColor("#111827"),
                            ),
                        )
                        for j, cell in enumerate(row)
                    ]
                )
            table = Table(table_data, colWidths=widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6D4AFF")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D7DCE8")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(table)
        for screenshot in screenshots_for(section["title"]):
            image_path = SCREENSHOT_DIR / screenshot["file"]
            if image_path.exists():
                story.append(Spacer(1, 0.05 * inch))
                story.append(pdf_screenshot_flowable(image_path))
                story.append(Paragraph(screenshot["caption"], styles["CaptionES"]))
        story.append(Spacer(1, 0.06 * inch))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def write_markdown(path: Path) -> None:
    lines = [
        "# Manual de usuario - Consola VYVA x Cruz Roja",
        "",
        f"Versión: {VERSION}",
        f"Actualizado: {UPDATED_ES}",
        f"URL: {MANUAL_URL}",
        "",
    ]
    for section in SECTIONS:
        lines.append(f"## {section['title']}")
        lines.extend(section.get("paras", []))
        lines.extend([f"- {item}" for item in section.get("bullets", [])])
        lines.extend([f"{i}. {item}" for i, item in enumerate(section.get("steps", []), start=1)])
        if "table" in section:
            rows = section["table"]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        for screenshot in screenshots_for(section["title"]):
            lines.append("")
            lines.append(f"![{screenshot['caption']}](../current/screenshots-from-english/{screenshot['file']})")
            lines.append(f"*{screenshot['caption']}*")
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def update_manifest() -> None:
    metadata_path = PUBLIC_DIR / "manual-version.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    metadata["languages"] = {
        "en": {
            "title": "VYVA Admin Console User Manual",
            "latestPdf": "/manuals/VYVA_Admin_Console_User_Manual.pdf",
            "archivePdf": "/manuals/archive/VYVA_Admin_Console_User_Manual_2026-06-20.pdf",
        },
        "es": {
            "title": "Manual de usuario - Consola VYVA x Cruz Roja",
            "latestPdf": "/manuals/VYVA_Admin_Console_User_Manual_ES.pdf",
            "archivePdf": "/manuals/archive/VYVA_Admin_Console_User_Manual_ES_2026-06-20.pdf",
        },
    }
    metadata["latestSpanishPdf"] = "/manuals/VYVA_Admin_Console_User_Manual_ES.pdf"
    metadata["spanishManualUrl"] = MANUAL_URL
    metadata_path.write_text(json.dumps(metadata, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def main() -> None:
    ensure_dirs()
    extract_screenshots()
    docx_root = ROOT / f"{OUT_NAME}.docx"
    pdf_root = ROOT / f"{OUT_NAME}.pdf"
    build_docx(docx_root)
    build_pdf(pdf_root)

    targets = [
        CURRENT_DIR / f"{OUT_NAME}.docx",
        VERSION_DIR / f"{OUT_NAME}.docx",
    ]
    for target in targets:
        shutil.copy2(docx_root, target)

    pdf_targets = [
        CURRENT_DIR / f"{OUT_NAME}.pdf",
        VERSION_DIR / f"{OUT_NAME}.pdf",
        PUBLIC_DIR / f"{OUT_NAME}.pdf",
        ARCHIVE_DIR / f"{OUT_NAME}_{VERSION}.pdf",
    ]
    for target in pdf_targets:
        shutil.copy2(pdf_root, target)

    write_markdown(SOURCE_DIR / "manual-spanish-2026-06-20.md")
    update_manifest()
    print(f"created {docx_root}")
    print(f"created {pdf_root}")


if __name__ == "__main__":
    main()
